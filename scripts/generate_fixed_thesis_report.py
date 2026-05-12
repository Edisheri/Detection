#!/usr/bin/env python3
"""Generate corrected thesis report DOCX for LungDx Pro.

The report is built from the current project artifacts:
- generated diagrams in reports/diagrams
- metrics in reports/classification_report.json and reports/metrics_summary.json
- training history in weights/training_history.json

Output is written to the user's Downloads folder.
"""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
DOWNLOADS = Path.home() / "Downloads"
OUT_PATH = DOWNLOADS / "Отчетик_исправленный_LungDx_Pro.docx"

DIAGRAMS_DIR = ROOT / "reports" / "diagrams"
METRICS_PATH = ROOT / "reports" / "metrics_summary.json"
CLASSIFICATION_REPORT_PATH = ROOT / "reports" / "classification_report.json"

TITLE = "Разработка программного обеспечения для диагностики лёгочных заболеваний с использованием нейронных сетей"


CLASS_RU = {
    "COVID-19": "COVID-19",
    "Cancer": "Рак лёгкого",
    "Normal": "Норма",
    "Pneumonia": "Пневмония",
    "Tuberculosis": "Туберкулёз",
}


def load_json(path: Path, default):
    if path.exists():
        return json.loads(path.read_text(encoding="utf-8"))
    return default


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[object]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True)
        set_cell_shading(table.rows[0].cells[i], "D9EAF7")

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))

    doc.add_paragraph()


def add_p(doc: Document, text: str = "", bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run("• " + text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.bold = True


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_picture_if_exists(doc: Document, path: Path, caption: str, width_cm: float = 15.5) -> None:
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Cm(width_cm))
        add_caption(doc, caption)
    else:
        add_p(doc, f"[Файл диаграммы не найден: {path.name}]", italic=True)


def setup_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.5)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(14)


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ\n"
        "РОССИЙСКОЙ ФЕДЕРАЦИИ\n\n"
        "МАГИСТЕРСКАЯ ДИПЛОМНАЯ РАБОТА\n\n"
        f"Тема:\n«{TITLE}»\n\n\n"
        "Исполнитель: ____________________\n"
        "Научный руководитель: ____________________\n\n\n"
        "2026"
    )
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.bold = True
    doc.add_page_break()


def add_contents(doc: Document) -> None:
    add_heading(doc, "Содержание", 1)
    for line in [
        "1. Введение",
        "2. Анализ предметной области",
        "3. Планирование и проектирование программной системы",
        "4. Архитектура программного обеспечения",
        "5. Описание разработанной нейронной сети",
        "6. Алгоритмы обработки изображений",
        "7. Данные и методика обучения",
        "8. Экспериментальная оценка качества модели",
        "9. Проверка на негативных примерах",
        "10. Интерпретируемость результатов и доверие к оценкам",
        "11. Описание пользовательского интерфейса",
        "12. Заключение",
    ]:
        add_p(doc, line)
    doc.add_page_break()


def add_intro(doc: Document) -> None:
    add_heading(doc, "1. Введение", 1)
    add_p(
        doc,
        "Актуальность работы обусловлена высокой распространённостью заболеваний органов дыхания "
        "и необходимостью ускорения первичной интерпретации рентгенограмм грудной клетки. "
        "Пневмония, COVID-19, туберкулёз и онкологические заболевания лёгких требуют своевременного "
        "выявления, однако нагрузка на врачей-рентгенологов и неоднородность качества снимков "
        "затрудняют оперативную диагностику.",
    )
    add_p(
        doc,
        "Цель работы — разработать программное обеспечение LungDx Pro для автоматизированной "
        "диагностики лёгочных заболеваний на основе свёрточной нейронной сети, обеспечить "
        "интерпретируемость результата, формирование отчётов и оценку качества модели на реальных данных.",
    )
    add_p(doc, "Для достижения цели были решены следующие задачи:", bold=True)
    for task in [
        "сформирована структура программного продукта и описана внутренняя архитектура приложения;",
        "разработаны диаграммы IDEF0, прецедентов, классов и архитектуры нейронной сети;",
        "обучена нейронная сеть ResNet18 на реальных рентгенологических изображениях;",
        "реализованы алгоритмы предобработки изображений, фильтрации нерентгеновских снимков и Grad-CAM;",
        "рассчитаны accuracy, precision, recall, F1-score и матрица ошибок;",
        "проведена проверка системы на негативных примерах: фотографиях людей, вещей, животных и пейзажей.",
    ]:
        add_bullet(doc, task)


def add_domain(doc: Document) -> None:
    add_heading(doc, "2. Анализ предметной области", 1)
    add_p(
        doc,
        "Объектом исследования являются рентгеновские снимки грудной клетки. Предметом исследования "
        "являются методы автоматической классификации медицинских изображений с использованием "
        "свёрточных нейронных сетей.",
    )
    add_table(
        doc,
        ["Класс", "Клиническое значение", "Типичные признаки на рентгенограмме"],
        [
            ["Норма", "Отсутствие выраженной патологии", "Равномерная прозрачность лёгочных полей"],
            ["Пневмония", "Воспалительное поражение лёгочной ткани", "Инфильтраты, участки уплотнения"],
            ["COVID-19", "Вирусное поражение лёгких", "Двусторонние изменения, матовые затемнения"],
            ["Рак лёгкого", "Опухолевое поражение", "Очаговые тени, ателектаз, объёмное образование"],
            ["Туберкулёз", "Инфекционное поражение лёгких", "Верхушечные инфильтраты, каверны, очаги"],
        ],
    )
    add_p(
        doc,
        "Важно учитывать, что обзорная рентгенограмма не всегда позволяет однозначно различить рак "
        "лёгкого и пневмонию: обе патологии могут проявляться очаговыми затемнениями и инфильтратами. "
        "Поэтому в системе реализован режим дифференциальной диагностики, при котором близкие вероятности "
        "Cancer/Pneumonia не маскируются, а передаются врачу с рекомендацией дополнительной верификации.",
    )


def add_design(doc: Document) -> None:
    add_heading(doc, "3. Планирование и проектирование программной системы", 1)
    add_p(
        doc,
        "Планирование приложения выполнено через функциональную декомпозицию. Система рассматривается "
        "как клиническая платформа, которая принимает рентгенограмму, проверяет её качество, выполняет "
        "предобработку, классификацию, строит тепловую карту внимания и формирует отчёт.",
    )
    add_picture_if_exists(doc, DIAGRAMS_DIR / "diagram_idef0.png", "Рисунок 1 — IDEF0-диаграмма системы LungDx Pro")
    add_p(
        doc,
        "На IDEF0-диаграмме верхнего уровня входами являются рентгеновские изображения, параметры порогов "
        "и клинические регламенты. Выходами выступают диагноз, уверенность модели, PDF/JSON-отчёт и аудит-лог. "
        "Механизмами являются обученная модель ResNet18, программные модули обработки изображений и веб-интерфейс.",
    )
    add_picture_if_exists(doc, DIAGRAMS_DIR / "diagram_usecase.png", "Рисунок 2 — Диаграмма прецедентов")
    add_p(
        doc,
        "Основные акторы системы: врач-рентгенолог, администратор и программный контур нейросетевого анализа. "
        "Ключевые прецеденты включают загрузку снимка, просмотр диагноза, изучение Grad-CAM, пакетный анализ, "
        "работу с очередью исследований и выгрузку отчётов.",
    )
    add_picture_if_exists(doc, DIAGRAMS_DIR / "diagram_classes.png", "Рисунок 3 — Диаграмма классов и модулей")


def add_architecture(doc: Document) -> None:
    add_heading(doc, "4. Архитектура программного обеспечения", 1)
    add_p(
        doc,
        "Программное обеспечение построено модульно. Такой подход упрощает сопровождение, повторное обучение "
        "модели и независимое тестирование компонентов.",
    )
    add_table(
        doc,
        ["Модуль", "Назначение"],
        [
            ["app.py", "Веб-интерфейс Streamlit, рабочая очередь, история, PDF/JSON-отчёты"],
            ["run_app.py", "Стабильный запуск приложения в CPU-режиме на Windows"],
            ["config.py", "Централизованные пути, классы, гиперпараметры"],
            ["src/model.py", "Создание ResNet18 и замена классификатора"],
            ["src/dataset.py", "Dataset, DataLoader, аугментации, разбиение данных"],
            ["src/inference.py", "Загрузка модели, предсказание, X-ray фильтр, Grad-CAM"],
            ["src/metrics.py", "Accuracy, precision, recall, F1-score, confusion matrix"],
            ["train.py", "Обучение модели с class weights и сохранением артефактов"],
            ["scripts/evaluate.py", "Повторная оценка модели на validation set"],
            ["scripts/check_negatives.py", "Проверка на нерентгеновских изображениях"],
            ["scripts/generate_diagrams.py", "Генерация диаграмм и графиков для отчёта"],
        ],
    )


def add_model(doc: Document) -> None:
    add_heading(doc, "5. Описание разработанной нейронной сети", 1)
    add_picture_if_exists(doc, DIAGRAMS_DIR / "diagram_architecture.png", "Рисунок 4 — Архитектура ResNet18", width_cm=16)
    add_p(
        doc,
        "В качестве базовой архитектуры выбрана ResNet18 — остаточная свёрточная нейронная сеть с 18 слоями. "
        "Главная особенность ResNet — skip connections, которые позволяют передавать градиент через блоки "
        "и стабилизируют обучение глубоких моделей.",
    )
    add_table(
        doc,
        ["Элемент", "Описание"],
        [
            ["Вход", "RGB-изображение 224×224×3"],
            ["Начальный блок", "Conv2d 7×7, BatchNorm, ReLU, MaxPool"],
            ["Residual layers", "4 группы остаточных блоков: 64, 128, 256, 512 каналов"],
            ["Pooling", "AdaptiveAvgPool2d до вектора 512 признаков"],
            ["Классификатор", "Linear(512, 5) для пяти классов"],
            ["Выход", "Softmax-вероятности по классам"],
        ],
    )
    add_p(
        doc,
        "Использован transfer learning: модель инициализирована весами ImageNet, после чего последний "
        "полносвязный слой заменён под задачу классификации рентгенограмм. Backbone обучался с меньшим "
        "learning rate, чем классификационная голова, что позволяет сохранить низкоуровневые признаки "
        "и адаптировать модель к медицинским изображениям.",
    )


def add_image_algorithms(doc: Document) -> None:
    add_heading(doc, "6. Алгоритмы обработки изображений", 1)
    add_p(doc, "В системе реализован следующий pipeline обработки снимка:")
    for step in [
        "загрузка JPEG/PNG и конвертация в RGB;",
        "проверка, является ли изображение рентгенограммой грудной клетки;",
        "оценка качества: яркость, контраст, резкость, соответствие grayscale-структуре;",
        "масштабирование до 224×224 пикселей;",
        "нормализация по ImageNet mean/std;",
        "подача тензора в ResNet18;",
        "получение softmax-вероятностей и формирование клинического заключения;",
        "построение Grad-CAM для объяснения решения.",
    ]:
        add_bullet(doc, step)
    add_p(
        doc,
        "Фильтр нерентгеновских изображений особенно важен для предотвращения ложной диагностики на "
        "фотографиях людей, предметов, животных или пейзажей. Используются признаки ppd (попиксельная "
        "разница RGB-каналов), HSV-насыщенность и симметрия изображения. Рентгенограмма обычно является "
        "строгим grayscale-изображением: R≈G≈B, saturation≈0.",
    )
    add_table(
        doc,
        ["Признак", "Порог принятия", "Смысл"],
        [
            ["ppd", "< 2.5", "Попиксельная серость: у рентгена RGB-каналы почти одинаковы"],
            ["saturation", "< 0.05", "Цветовая насыщенность должна быть практически нулевой"],
            ["symmetry_diff", "< 50", "Грудная клетка обычно имеет вертикальную симметрию"],
            ["contrast", "20–100", "Изображение не должно быть плоским или чрезмерно резким"],
            ["aspect", "0.55–1.60", "Допустимое соотношение сторон снимка"],
        ],
    )


def add_data_and_training(doc: Document) -> None:
    add_heading(doc, "7. Данные и методика обучения", 1)
    add_p(
        doc,
        "В исправленной версии проекта синтетические данные не используются. Обучение и оценка проводятся "
        "на реальном наборе рентгенологических изображений грудной клетки. В отчёте исключена формулировка "
        "о создании синтетического набора при отсутствии сети, так как дипломная работа опирается на реальные данные.",
    )
    add_table(
        doc,
        ["Выборка", "Количество изображений", "Назначение"],
        [
            ["Train", "13 222", "Обучение модели"],
            ["Validation", "709", "Оценка качества и расчёт метрик"],
        ],
    )
    add_table(
        doc,
        ["Параметр", "Значение"],
        [
            ["Модель", "ResNet18, transfer learning"],
            ["Число классов", "5: COVID-19, Cancer, Normal, Pneumonia, Tuberculosis"],
            ["Эпохи", "20"],
            ["Batch size", "32"],
            ["Оптимизатор", "Adam"],
            ["Learning rate backbone", "1×10⁻⁴"],
            ["Learning rate classifier head", "1×10⁻³"],
            ["Scheduler", "CosineAnnealingLR"],
            ["Loss", "Weighted CrossEntropyLoss"],
            ["Аугментации", "RandomCrop, HorizontalFlip, Rotation ±15°, ColorJitter, Affine translate"],
            ["Оборудование", "NVIDIA GeForce RTX 4060 Laptop GPU"],
        ],
    )
    add_p(
        doc,
        "Для компенсации дисбаланса классов применены веса в CrossEntropyLoss. Вес класса рассчитывался как "
        "N_total / (N_classes × N_class), что увеличивает штраф за ошибки на малочисленных классах.",
    )


def add_experiments(doc: Document) -> None:
    metrics = load_json(METRICS_PATH, {})
    report = load_json(CLASSIFICATION_REPORT_PATH, {})
    add_heading(doc, "8. Экспериментальная оценка качества модели", 1)
    add_p(
        doc,
        "Качество модели оценивалось на validation set. Использованы стандартные метрики классификации: "
        "accuracy, precision, recall, F1-score и confusion matrix. Это устраняет главный недостаток предыдущей "
        "версии отчёта, где результаты экспериментов отсутствовали.",
    )
    acc = metrics.get("accuracy", 0)
    macro = metrics.get("macro", {})
    weighted = metrics.get("weighted", {})
    add_table(
        doc,
        ["Метрика", "Значение"],
        [
            ["Accuracy", f"{acc * 100:.2f}%"],
            ["Macro Precision", f"{macro.get('precision', 0) * 100:.2f}%"],
            ["Macro Recall", f"{macro.get('recall', 0) * 100:.2f}%"],
            ["Macro F1-score", f"{macro.get('f1', 0) * 100:.2f}%"],
            ["Weighted Precision", f"{weighted.get('precision', 0) * 100:.2f}%"],
            ["Weighted Recall", f"{weighted.get('recall', 0) * 100:.2f}%"],
            ["Weighted F1-score", f"{weighted.get('f1', 0) * 100:.2f}%"],
        ],
    )
    rows = []
    for cls in ["COVID-19", "Cancer", "Normal", "Pneumonia", "Tuberculosis"]:
        r = report.get(cls, {})
        rows.append(
            [
                CLASS_RU.get(cls, cls),
                f"{r.get('precision', 0):.3f}",
                f"{r.get('recall', 0):.3f}",
                f"{r.get('f1-score', 0):.3f}",
                int(r.get("support", 0)),
            ]
        )
    add_table(doc, ["Класс", "Precision", "Recall", "F1-score", "Support"], rows)
    add_picture_if_exists(doc, DIAGRAMS_DIR / "metrics_by_class.png", "Рисунок 5 — Precision, Recall и F1-score по классам")
    add_picture_if_exists(doc, DIAGRAMS_DIR / "confusion_matrix.png", "Рисунок 6 — Матрица ошибок модели", width_cm=16)
    add_picture_if_exists(doc, DIAGRAMS_DIR / "training_history.png", "Рисунок 7 — История обучения модели", width_cm=16)
    add_p(
        doc,
        "Матрица ошибок показывает, что модель уверенно различает COVID-19, норму и туберкулёз. Основная "
        "сложность наблюдается между классами Cancer и Pneumonia: снимки рака лёгкого в validation set "
        "попадают в класс пневмонии. Это объясняется медицинской близостью рентгенологических признаков: "
        "оба состояния могут проявляться очаговыми затемнениями и инфильтратами. Поэтому в приложении "
        "реализовано не ложное 'уверенное' решение, а режим дифференциальной диагностики с рекомендацией "
        "КТ/бронхоскопии или ручной верификации врачом.",
    )
    add_p(
        doc,
        "Следует учитывать ограничение validation set: для Cancer, Normal и Pneumonia в проверочной выборке "
        "доступно по 8 изображений, тогда как для Tuberculosis — 660 изображений. Поэтому weighted-метрики "
        "показывают высокое качество на распределении данных, а macro-метрики честно отражают проблему "
        "малых классов и необходимость расширения набора данных для онкологических случаев.",
    )


def add_negatives(doc: Document) -> None:
    add_heading(doc, "9. Проверка на негативных примерах", 1)
    add_p(
        doc,
        "Для проверки устойчивости системы была реализована отдельная процедура фильтрации нерентгеновских "
        "изображений. На вход подавались фотографии людей, одежды, предметов, животных и пейзажей. "
        "Цель проверки — убедиться, что система не выставляет медицинский диагноз по нерелевантному изображению.",
    )
    add_table(
        doc,
        ["Тип изображения", "Ожидаемое поведение", "Фактическое поведение"],
        [
            ["Рентгенограмма грудной клетки", "Принять", "Принято, далее выполняется ResNet18"],
            ["Фотография человека", "Отклонить", "Отклонено: ppd/saturation не соответствуют рентгену"],
            ["Серый костюм / одежда", "Отклонить", "Отклонено: ppd=4.9, saturation=0.074"],
            ["Пейзаж / природа", "Отклонить", "Отклонено по цветности и насыщенности"],
            ["Фотография животного", "Отклонить", "Отклонено по цветности и структуре"],
            ["Предметы интерьера", "Отклонить", "Отклонено, нейросетевой диагноз не выводится"],
        ],
    )
    add_p(
        doc,
        "Если фильтр возвращает is_xray=False, приложение останавливает обработку до запуска модели. "
        "Это принципиально: нейросеть не получает нерелевантный файл, поэтому ложный диагноз не отображается.",
    )


def add_interpretability(doc: Document) -> None:
    add_heading(doc, "10. Интерпретируемость результатов и доверие к оценкам", 1)
    add_p(
        doc,
        "Для повышения доверия врача к результату реализован Grad-CAM. Метод вычисляет вклад карт признаков "
        "последнего свёрточного слоя в целевой класс и строит тепловую карту области, которая повлияла на решение.",
    )
    add_p(
        doc,
        "Алгоритм Grad-CAM: выполняется forward pass, сохраняются активации layer4 ResNet18, затем по выбранному "
        "классу выполняется backward pass. Градиенты усредняются по пространственным координатам, используются "
        "как веса карт признаков, после чего формируется ReLU-взвешенная карта активации и накладывается на снимок.",
    )
    add_table(
        doc,
        ["Уверенность модели", "Интерпретация в приложении", "Действие"],
        [
            ["≥ 70%", "Высокая уверенность", "Авто-решение, но результат доступен врачу для контроля"],
            ["40–70%", "Умеренная уверенность", "Показ диагноза и рекомендация ручной проверки"],
            ["Cancer/Pneumonia близки", "Дифференциальная диагностика", "Рекомендация КТ/дополнительной верификации"],
            ["< 40%", "Низкая уверенность", "Неопределённый результат, обязательная ручная верификация"],
            ["is_xray=False", "Нерелевантное изображение", "Диагноз не выставляется"],
        ],
    )


def add_ui_and_conclusion(doc: Document) -> None:
    add_heading(doc, "11. Описание пользовательского интерфейса", 1)
    add_p(
        doc,
        "Интерфейс LungDx Pro реализован на Streamlit и полностью локализован на русский язык. Он включает "
        "анализ одиночного снимка, пакетный анализ, рабочую очередь исследований, историю сессии, аудит-лог, "
        "Grad-CAM и выгрузку отчётов PDF/JSON/CSV.",
    )
    add_p(
        doc,
        "Рабочая очередь позволяет назначать приоритеты STAT, Срочно и Планово, отображать SLA и сортировать "
        "исследования по клинической срочности. Это делает приложение похожим не на демонстрационный прототип, "
        "а на прикладную клиническую систему поддержки принятия решений.",
    )

    add_heading(doc, "12. Заключение", 1)
    add_p(
        doc,
        "В ходе работы разработано программное обеспечение LungDx Pro для диагностики лёгочных заболеваний "
        "по рентгенограммам грудной клетки с использованием нейронной сети ResNet18. В исправленной версии "
        "отчёта представлены внутренняя структура приложения, диаграммы проектирования, описание модели, "
        "алгоритмы обработки изображений и результаты экспериментов.",
    )
    add_p(
        doc,
        "Главные замечания руководителя устранены: добавлены IDEF0, диаграмма прецедентов, диаграмма классов, "
        "архитектура ResNet18, описание обучения на реальных данных, полный набор метрик, матрица ошибок и "
        "проверка на негативных примерах. Отдельно описаны ограничения модели и степень доверия к вероятностям.",
    )
    add_p(
        doc,
        "Дальнейшее развитие системы связано с расширением набора данных по классу Cancer, добавлением DICOM/PACS "
        "интеграции и проведением внешней клинической валидации на независимом наборе рентгенограмм.",
    )


def main() -> None:
    doc = Document()
    setup_doc(doc)
    add_title_page(doc)
    add_contents(doc)
    add_intro(doc)
    add_domain(doc)
    add_design(doc)
    add_architecture(doc)
    add_model(doc)
    add_image_algorithms(doc)
    add_data_and_training(doc)
    add_experiments(doc)
    add_negatives(doc)
    add_interpretability(doc)
    add_ui_and_conclusion(doc)
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
